from docx import Document
from datetime import datetime
from project.settings import MEDIA_ROOT
from home.planejamento.api_ia.main_ia import generate_planning_ia
import locale
from home.utils.unique_slugify import generate_slug

locale.setlocale(locale.LC_TIME, "pt_BR.utf8")

DEFAULT_SAVE_FOLDER = MEDIA_ROOT / 'files_docx_generated'

def init_generate_document(matters: list, date: str, term_for_ia: dict, extra: str, basedSep: bool, school, teacher):
    try:
        ids = 0
        date_formated = datetime.strptime(date, '%Y-%m-%d')
        date_exibs = datetime.strftime(date_formated, '%d de %B de %Y')
        
        document = Document()
        style = document.styles['Normal']
        style.paragraph_format.line_spacing = 1
        font = style.font
        font.name = 'Arial'

        # verifica a escolha do usuário e adicona no cabeçalho o que ele quiser
        if school and teacher:
            document.add_heading(f'{school}, {date_exibs} - {teacher.first_name} {teacher.last_name}', 0)
        elif school:
            document.add_heading(f'{school}, {date_exibs}', 0)
        elif teacher:
            document.add_heading(f'{date_exibs} - {teacher.first_name} {teacher.last_name}', 0)
        else:
            document.add_heading(f'{date_exibs}', 0)


        for matter in matters:
            if basedSep:
                hour_formated = datetime.strptime(str(matter.hour), '%H:%M:%S')
                sep_exibs = datetime.strftime(hour_formated, '%Hh%M')
            else:
                ids += 1
                sep_exibs = f'{ids}ª Aula'

            info_aula = document.add_paragraph(f'{sep_exibs} - ')
            info_aula.add_run(matter.matter).bold = True

            term = term_for_ia[f'{matter.pk}']
            if term:
                document.add_paragraph(generate_planning_ia(term))
            else:
                document.add_paragraph('').add_run('(escreva aqui)').italic = True
            
        if extra:
            document.add_paragraph(extra)

        slug_name = generate_slug(15)
        file_name = DEFAULT_SAVE_FOLDER / f'planejamento_{slug_name}.docx'

        document.save(file_name)
        return slug_name
    except Exception as e:
        # if raise any error, return False (this is treated on view)
        print(e)
        return False